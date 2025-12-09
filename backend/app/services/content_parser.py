import re
from typing import Optional
import PyPDF2
import docx
from bs4 import BeautifulSoup
import markdown
import logging

logger = logging.getLogger(__name__)


class ContentParser:
    """Parse various content formats into plain text."""

    @staticmethod
    async def parse_content(file_path: str, content_type: str) -> str:
        """Parse content based on type."""
        try:
            if content_type == "html":
                return await ContentParser._parse_html(file_path)
            elif content_type == "markdown":
                return await ContentParser._parse_markdown(file_path)
            elif content_type == "pdf":
                return await ContentParser._parse_pdf(file_path)
            elif content_type == "docx":
                return await ContentParser._parse_docx(file_path)
            else:
                raise ValueError(f"Unsupported content type: {content_type}")

        except Exception as e:
            logger.error(f"Error parsing {content_type}: {str(e)}")
            raise

    @staticmethod
    async def _parse_html(file_path: str) -> str:
        """Parse HTML file to plain text."""
        with open(file_path, 'r', encoding='utf-8') as f:
            html_content = f.read()

        soup = BeautifulSoup(html_content, 'html.parser')

        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()

        # Get text
        text = soup.get_text()

        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = ' '.join(chunk for chunk in chunks if chunk)

        return text

    @staticmethod
    async def _parse_markdown(file_path: str) -> str:
        """Parse Markdown file to plain text."""
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # Convert to HTML first, then to text
        html = markdown.markdown(md_content)
        soup = BeautifulSoup(html, 'html.parser')
        text = soup.get_text()

        return text.strip()

    @staticmethod
    async def _parse_pdf(file_path: str) -> str:
        """Parse PDF file to plain text."""
        text = []

        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)

            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text.append(page.extract_text())

        return '\n'.join(text)

    @staticmethod
    async def _parse_docx(file_path: str) -> str:
        """Parse DOCX file to plain text."""
        doc = docx.Document(file_path)

        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)

        return '\n'.join(text)


class ContentParserService:
    """
    Synchronous wrapper for ContentParser.
    Used by PreprocessingService which needs sync methods.
    """
    
    def parse_pdf(self, file_path: str) -> str:
        """Synchronously parse PDF file."""
        text = []
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text.append(page.extract_text())
        return '\n'.join(text)
    
    def parse_docx(self, file_path: str) -> str:
        """Synchronously parse DOCX file."""
        doc = docx.Document(file_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return '\n'.join(text)


content_parser = ContentParser()
